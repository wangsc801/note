from docx import Document


def get_table_content(fin):
    document = Document(fin)
    result = []
    for table in document.tables:
        cur_table = []
        for row in table.rows:
            cur_row = []
            for cell in row.cells:
                cur_row.append(cell.text)
            cur_table.append(cur_row)
        result.append(cur_table)
    return result


def get_table_content_of_row(fin, row_num):
    document = Document(fin)
    result = []
    for tb in document.tables:
        for cell in tb.row_cells(row_num):
            result.append(cell.text)
    return result
